import os
import re
from datetime import datetime, timezone

from backend.core.logger import logger

# Try imports for optional document generation libraries to allow loading even
# if missing,
# but we will raise appropriate errors if they are called and not available.
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def validate_inputs(report: dict, output_path: str):
    """
    Validates the structure of the input report dictionary and the output path.
    """
    if not isinstance(report, dict):
        raise ValueError("Report input must be a dictionary")

    if not output_path or not isinstance(output_path, str):
        raise ValueError("Output path must be a non-empty string")

    # Check if target directory exists
    dir_name = os.path.dirname(output_path)
    if dir_name and not os.path.isdir(dir_name):
        raise ValueError(f"Target directory does not exist: {dir_name}")

    # Check if output path is a directory
    if os.path.isdir(output_path):
        raise ValueError(f"Output path is a directory, not a file: {output_path}")

    required_fields = ["topic", "title", "final_markdown", "review", "changes_applied"]
    for field in required_fields:
        if field not in report:
            raise ValueError(f"Missing required field: {field}")

    review = report["review"]
    if not isinstance(review, dict):
        raise ValueError("Field 'review' must be a dictionary")

    if not isinstance(report["changes_applied"], list):
        raise ValueError("Field 'changes_applied' must be a list")


def format_inline_markdown(text: str) -> str:
    """
    Converts inline Markdown formatting (bold, italic, code) to ReportLab
    Paragraph HTML tags.
    """
    # Normalize underscores to asterisks for simple markdown parsing
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # Escape sequences normalization
    text = re.sub(r"___(.*?)___", r"***\1***", text)
    text = re.sub(r"__(.*?)__", r"**\1**", text)
    text = re.sub(r"_(.*?)_", r"*\1*", text)

    # Bold-Italic: ***text*** -> <b><i>text</i></b>
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"<b><i>\1</i></b>", text)

    # Bold: **text** -> <b>text</b>
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)

    # Italic: *text* -> <i>text</i>
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)

    # Inline code: `code` -> <font face="Courier">code</font>
    text = re.sub(r"`(.*?)`", r'<font face="Courier">\1</font>', text)

    return text


def add_formatted_text_to_paragraph(text: str, paragraph):
    """
    Parses inline markdown format and adds runs with styles to a python-docx paragraph.
    """
    # Normalize underscores to asterisks
    text = re.sub(r"___(.*?)___", r"***\1***", text)
    text = re.sub(r"__(.*?)__", r"**\1**", text)
    text = re.sub(r"_(.*?)_", r"*\1*", text)

    pattern = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def parse_markdown_to_flowables(markdown_text: str, styles) -> list:
    """
    Parses Markdown text and returns a list of ReportLab Flowables.
    """
    flowables = []

    # Split by code blocks first
    lines = markdown_text.splitlines()
    blocks = []
    in_code_block = False
    current_block = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                blocks.append(("code", "\n".join(current_block)))
                current_block = []
            else:
                if current_block:
                    blocks.append(("text", "\n".join(current_block)))
                    current_block = []
                in_code_block = True
        else:
            current_block.append(line)

    if current_block:
        blocks.append(("code" if in_code_block else "text", "\n".join(current_block)))

    for block_type, content in blocks:
        if block_type == "code":
            escaped_code = (
                content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            code_style = ParagraphStyle(
                "CodeBlock",
                parent=styles["Normal"],
                fontName="Courier",
                fontSize=9,
                leading=11,
                textColor=colors.HexColor("#2d3748"),
                backColor=colors.HexColor("#f7fafc"),
                borderColor=colors.HexColor("#e2e8f0"),
                borderWidth=1,
                borderPadding=8,
                spaceAfter=12,
            )
            flowables.append(Paragraph(escaped_code.replace("\n", "<br/>"), code_style))
            flowables.append(Spacer(1, 6))
        else:
            paragraphs = re.split(r"\n\s*\n", content)
            for p_text in paragraphs:
                p_text = p_text.strip()
                if not p_text:
                    continue

                # Check for headings
                if p_text.startswith("# "):
                    h_text = p_text[2:].strip()
                    flowables.append(Paragraph(h_text, styles["Heading1"]))
                    flowables.append(Spacer(1, 8))
                elif p_text.startswith("## "):
                    h_text = p_text[3:].strip()
                    flowables.append(Paragraph(h_text, styles["Heading2"]))
                    flowables.append(Spacer(1, 6))
                elif p_text.startswith("### "):
                    h_text = p_text[4:].strip()
                    flowables.append(Paragraph(h_text, styles["Heading3"]))
                    flowables.append(Spacer(1, 6))
                elif p_text.startswith("- ") or p_text.startswith("* "):
                    items = p_text.split("\n")
                    for item in items:
                        item = item.strip()
                        if item.startswith("- ") or item.startswith("* "):
                            item_content = item[2:].strip()
                            formatted_item = format_inline_markdown(item_content)
                            flowables.append(
                                Paragraph(f"&bull; {formatted_item}", styles["Bullet"])
                            )
                    flowables.append(Spacer(1, 6))
                elif p_text.startswith("1. ") or re.match(r"^\d+\.\s", p_text):
                    items = p_text.split("\n")
                    for item in items:
                        item = item.strip()
                        match = re.match(r"^(\d+)\.\s(.*)", item)
                        if match:
                            num = match.group(1)
                            item_content = match.group(2).strip()
                            formatted_item = format_inline_markdown(item_content)
                            flowables.append(
                                Paragraph(f"{num}. {formatted_item}", styles["Normal"])
                            )
                    flowables.append(Spacer(1, 6))
                else:
                    formatted_p = format_inline_markdown(p_text)
                    flowables.append(Paragraph(formatted_p, styles["BodyText"]))
                    flowables.append(Spacer(1, 8))

    return flowables


def parse_markdown_to_docx(markdown_text: str, doc):
    """
    Parses Markdown text and appends corresponding paragraphs/headers/lists to
    a python-docx Document.
    """
    lines = markdown_text.splitlines()
    blocks = []
    in_code_block = False
    current_block = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                blocks.append(("code", "\n".join(current_block)))
                current_block = []
            else:
                if current_block:
                    blocks.append(("text", "\n".join(current_block)))
                    current_block = []
                in_code_block = True
        else:
            current_block.append(line)

    if current_block:
        blocks.append(("code" if in_code_block else "text", "\n".join(current_block)))

    for block_type, content in blocks:
        if block_type == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        else:
            paragraphs = re.split(r"\n\s*\n", content)
            for p_text in paragraphs:
                p_text = p_text.strip()
                if not p_text:
                    continue

                if p_text.startswith("# "):
                    doc.add_heading(p_text[2:].strip(), level=1)
                elif p_text.startswith("## "):
                    doc.add_heading(p_text[3:].strip(), level=2)
                elif p_text.startswith("### "):
                    doc.add_heading(p_text[4:].strip(), level=3)
                elif p_text.startswith("- ") or p_text.startswith("* "):
                    items = p_text.split("\n")
                    for item in items:
                        item = item.strip()
                        if item.startswith("- ") or item.startswith("* "):
                            p = doc.add_paragraph(style="List Bullet")
                            add_formatted_text_to_paragraph(item[2:].strip(), p)
                elif p_text.startswith("1. ") or re.match(r"^\d+\.\s", p_text):
                    items = p_text.split("\n")
                    for item in items:
                        item = item.strip()
                        match = re.match(r"^(\d+)\.\s(.*)", item)
                        if match:
                            item_content = match.group(2).strip()
                            p = doc.add_paragraph(style="List Number")
                            add_formatted_text_to_paragraph(item_content, p)
                else:
                    p = doc.add_paragraph()
                    add_formatted_text_to_paragraph(p_text, p)


# Module-level implementation functions
def export_markdown(report: dict, output_path: str) -> str:
    logger.info(f"Export started - format: markdown, path: {output_path}")
    try:
        validate_inputs(report, output_path)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(report["final_markdown"])

        logger.info(f"Export completed - format: markdown, path: {output_path}")
        return os.path.abspath(output_path)
    except Exception as e:
        logger.error(
            f"Export failed - format: markdown, path: {output_path}, error: {str(e)}"
        )
        raise


def export_pdf(report: dict, output_path: str) -> str:
    logger.info(f"Export started - format: pdf, path: {output_path}")
    try:
        validate_inputs(report, output_path)

        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("ReportLab is not installed or available.")

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()

        # Modify existing styles in-place to avoid namespace collisions and
        # ensure clean output
        styles["Title"].fontName = "Helvetica-Bold"
        styles["Title"].fontSize = 24
        styles["Title"].leading = 28
        styles["Title"].spaceAfter = 15
        styles["Title"].textColor = colors.HexColor("#1a365d")

        styles["BodyText"].fontName = "Helvetica"
        styles["BodyText"].fontSize = 10.5
        styles["BodyText"].leading = 15
        styles["BodyText"].spaceAfter = 8
        styles["BodyText"].textColor = colors.HexColor("#2d3748")

        styles["Heading1"].fontName = "Helvetica-Bold"
        styles["Heading1"].fontSize = 16
        styles["Heading1"].leading = 20
        styles["Heading1"].spaceBefore = 14
        styles["Heading1"].spaceAfter = 8
        styles["Heading1"].textColor = colors.HexColor("#2b6cb0")
        styles["Heading1"].keepWithNext = True

        styles["Heading2"].fontName = "Helvetica-Bold"
        styles["Heading2"].fontSize = 13
        styles["Heading2"].leading = 17
        styles["Heading2"].spaceBefore = 12
        styles["Heading2"].spaceAfter = 6
        styles["Heading2"].textColor = colors.HexColor("#2d3748")
        styles["Heading2"].keepWithNext = True

        styles["Heading3"].fontName = "Helvetica-Bold"
        styles["Heading3"].fontSize = 11
        styles["Heading3"].leading = 14
        styles["Heading3"].spaceBefore = 10
        styles["Heading3"].spaceAfter = 4
        styles["Heading3"].textColor = colors.HexColor("#4a5568")
        styles["Heading3"].keepWithNext = True

        styles["Bullet"].fontName = "Helvetica"
        styles["Bullet"].fontSize = 10
        styles["Bullet"].leading = 14
        styles["Bullet"].leftIndent = 15
        styles["Bullet"].firstLineIndent = -10
        styles["Bullet"].spaceAfter = 4
        styles["Bullet"].textColor = colors.HexColor("#2d3748")

        flowables = []

        # 1. Document Title
        flowables.append(
            Paragraph(report.get("title", "Research Report"), styles["Title"])
        )
        flowables.append(Spacer(1, 10))

        # 2. Evaluation / Metadata Table
        review = report.get("review") or {}
        score = review.get("score", "N/A")
        strengths = review.get("strengths") or []
        issues = review.get("issues") or []
        suggestions = review.get("suggestions") or []
        changes_applied = report.get("changes_applied") or []
        timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

        metadata_title_style = ParagraphStyle(
            "MetadataTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#1a365d"),
        )

        label_style = ParagraphStyle(
            "MetadataLabel",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor("#4a5568"),
        )

        val_style = ParagraphStyle(
            "MetadataVal",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#2d3748"),
        )

        review_table_data = [
            [
                Paragraph(
                    "<b>Report Evaluation & Metadata Summary</b>", metadata_title_style
                ),
                "",
            ],
            [
                Paragraph("Topic", label_style),
                Paragraph(report.get("topic", ""), val_style),
            ],
            [
                Paragraph("Evaluation Score", label_style),
                Paragraph(f"<b>{score}/100</b>", val_style),
            ],
            [
                Paragraph("Strengths", label_style),
                Paragraph(
                    "<br/>".join(
                        [f"&bull; {format_inline_markdown(s)}" for s in strengths]
                    )
                    if strengths
                    else "None",
                    val_style,
                ),
            ],
            [
                Paragraph("Issues Identified", label_style),
                Paragraph(
                    "<br/>".join(
                        [f"&bull; {format_inline_markdown(i)}" for i in issues]
                    )
                    if issues
                    else "None",
                    val_style,
                ),
            ],
            [
                Paragraph("Suggestions", label_style),
                Paragraph(
                    "<br/>".join(
                        [f"&bull; {format_inline_markdown(s)}" for s in suggestions]
                    )
                    if suggestions
                    else "None",
                    val_style,
                ),
            ],
            [
                Paragraph("Changes Applied", label_style),
                Paragraph(
                    "<br/>".join(
                        [f"&bull; {format_inline_markdown(c)}" for c in changes_applied]
                    )
                    if changes_applied
                    else "None",
                    val_style,
                ),
            ],
            [
                Paragraph("Generated Timestamp", label_style),
                Paragraph(timestamp, val_style),
            ],
        ]

        t = Table(review_table_data, colWidths=[130, 374])
        t.setStyle(
            TableStyle(
                [
                    ("SPAN", (0, 0), (1, 0)),
                    ("BACKGROUND", (0, 0), (1, 0), colors.HexColor("#edf2f7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e0")),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )

        flowables.append(t)
        flowables.append(Spacer(1, 20))

        # 3. Render Markdown Content
        content_flowables = parse_markdown_to_flowables(
            report.get("final_markdown", ""), styles
        )
        flowables.extend(content_flowables)

        doc.build(flowables)

        logger.info(f"Export completed - format: pdf, path: {output_path}")
        return os.path.abspath(output_path)
    except Exception as e:
        logger.error(
            f"Export failed - format: pdf, path: {output_path}, error: {str(e)}"
        )
        raise


def export_docx(report: dict, output_path: str) -> str:
    logger.info(f"Export started - format: docx, path: {output_path}")
    try:
        validate_inputs(report, output_path)

        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed or available.")

        doc = Document()

        # 1. Title
        title_p = doc.add_paragraph()
        run = title_p.add_run(report.get("title", "Research Report"))
        run.bold = True
        run.font.size = Pt(24)
        run.font.name = "Arial"

        # 2. Evaluation Metadata Table
        review = report.get("review") or {}
        score = review.get("score", "N/A")
        strengths = review.get("strengths") or []
        issues = review.get("issues") or []
        suggestions = review.get("suggestions") or []
        changes_applied = report.get("changes_applied") or []
        timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

        # Insert a table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_meta_row(label, val):
            row = table.add_row()
            # Label
            cell_lbl = row.cells[0]
            cell_lbl.paragraphs[0].add_run(label).bold = True
            # Value
            cell_val = row.cells[1]
            p_val = cell_val.paragraphs[0]
            if isinstance(val, list):
                if val:
                    for i, item in enumerate(val):
                        if i > 0:
                            p_val = cell_val.add_paragraph()
                        add_formatted_text_to_paragraph(f"• {item}", p_val)
                else:
                    p_val.text = "None"
            else:
                p_val.text = str(val)

        # Main header row spanning both columns
        hdr_row = table.add_row()
        hdr_cell = hdr_row.cells[0]
        hdr_cell.merge(hdr_row.cells[1])
        hdr_run = hdr_cell.paragraphs[0].add_run("Report Evaluation & Metadata Summary")
        hdr_run.bold = True

        add_meta_row("Topic", report.get("topic", ""))
        add_meta_row("Evaluation Score", f"{score}/100")
        add_meta_row("Strengths", strengths)
        add_meta_row("Issues Identified", issues)
        add_meta_row("Suggestions", suggestions)
        add_meta_row("Changes Applied", changes_applied)
        add_meta_row("Generated Timestamp", timestamp)

        doc.add_paragraph()  # spacing

        # 3. Main content
        parse_markdown_to_docx(report.get("final_markdown", ""), doc)

        doc.save(output_path)

        logger.info(f"Export completed - format: docx, path: {output_path}")
        return os.path.abspath(output_path)
    except Exception as e:
        logger.error(
            f"Export failed - format: docx, path: {output_path}, error: {str(e)}"
        )
        raise


class ExportService:
    """
    ExportService provides methods to export research reports into multiple
    formats including Markdown, PDF, and DOCX.
    """

    @staticmethod
    def export_markdown(report: dict, output_path: str) -> str:
        return export_markdown(report, output_path)

    @staticmethod
    def export_pdf(report: dict, output_path: str) -> str:
        return export_pdf(report, output_path)

    @staticmethod
    def export_docx(report: dict, output_path: str) -> str:
        return export_docx(report, output_path)
